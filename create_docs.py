from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np

def create_doc():
    doc = Document()

    # Tiêu đề
    title = doc.add_heading('HƯỚNG DẪN TRIỂN KHAI VÀ DEPLOY HỆ THỐNG ĐẾM KHOAI TÂY (POTATO COUNTER)', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Mục lục hoặc lời nói đầu
    doc.add_paragraph('Tài liệu này được biên soạn cho nhóm 5 thành viên. Mỗi thành viên có thể dựa vào hướng dẫn này để setup từ môi trường cơ bản nhất đến lúc chạy được trên Máy tính cá nhân (Local) và Google Colab (Cloud/Deploy).')

    # PHẦN 1
    doc.add_heading('PHẦN 1: Chuẩn Bị Môi Trường (TẤT CẢ THÀNH VIÊN ĐỀU PHẢI LÀM)', level=1)
    p = doc.add_paragraph('Mỗi cá nhân muốn chạy được dự án này trên máy tính cá nhân cần phải Cài đặt môi trường Python.\n')
    run = p.add_run('Bước 1: ')
    run.bold = True
    p.add_run('Tải và cài đặt ngôn ngữ lập trình Python bản 3.9 đến 3.12 từ trang chủ python.org.\n')
    run = p.add_run('Bước 2: ')
    run.bold = True
    p.add_run('Sao chép (clone) hoặc tải thư mục mã nguồn dự án (Potato-Counter) về máy tính.\n')
    run = p.add_run('Bước 3: ')
    run.bold = True
    p.add_run('Mở Terminal (Command Prompt / PowerShell) lên, gõ lệnh cài đặt các thư viện cần thiết bằng lệnh sau:\n')
    doc.add_paragraph('pip install ultralytics opencv-python numpy scikit-image filterpy lap', style='Intense Quote')

    # PHẦN 2
    doc.add_heading('PHẦN 2: Triển khai trên Local - Dành cho (Thành viên 3 và 4)', level=1)
    p = doc.add_paragraph('Phần này hướng dẫn các bạn Thử nghiệm hệ thống, điều chỉnh giao diện, test lại vị trí đa giác đếm. Chỉ sử dụng CPU nên tốc độ sẽ chậm.\n')
    run = p.add_run('Cách chạy:\n')
    run.bold = True
    p.add_run('Mở Terminal tại thư mục Potato-Counter và gõ lệnh sau:\n')
    doc.add_paragraph('python manual_counter.py', style='Intense Quote')
    run = p.add_run('Cách vẽ lại vùng Polygon (Ví dụ đổi sang đếm băng chuyền bên phải):\n')
    run.bold = True
    p.add_run('1. Lấy khung hình 1280x720 của Video.\n2. Ném vào trang Roboflow (Universe) phần PolygonZone.\n3. Nhấn phím P để đồ theo khu vực cần đếm.\n4. Sao chép kết quả Tọa Độ numpy dán vào dòng 19 trong file manual_counter.py.\n')

    # PHẦN 3
    doc.add_heading('PHẦN 3: Deploy và Test Tốc độ Thực Tế - Google Colab (Dành cho Cả Nhóm/Thành viên 5)', level=1)
    p = doc.add_paragraph('Máy tính chạy CPU thường rất chậm (3-5 fps). Để hệ thống đạt chuẩn mượt mà 30+ fps, chúng ta cần Deploy (Đẩy mã nguồn) lên Máy Chủ mây Google Colab có trang bị Card Rời (GPU).\n')
    run = p.add_run('Bước 1: Truy cập Google Colab\n')
    run.bold = True
    p.add_run('1. Vào link: colab.research.google.com -> Tạo Notebook mới.\n')
    p.add_run('2. Vào menu Runtime (Thời gian chạy) -> Change runtime type (Đổi phiên bản) -> Chọn T4 GPU -> Nhấn Save.\n')
    run = p.add_run('\nBước 2: Tải các tệp quan trạng lên\n')
    run.bold = True
    p.add_run('Nhấn vào biểu tượng thư mục bên trái, tải 3 File cốt lõi này lên vùng trống:\n- Potato_model_best_weights.pt (Mạng Neural YOLO)\n- sort.py (Thuật toán Tracking)\n- ooo.mp4 (File Video băng chuyền)\n')
    run = p.add_run('\nBước 3: Dán Code Thực thi\n')
    run.bold = True
    p.add_run('Copy đoạn mã (Script) tích hợp sẵn dưới đây dán vào Ô lệnh của Colab và nhấn Nút (Run cell).')

    code_block = '''
!pip install ultralytics scikit-image filterpy lap
from ultralytics import YOLO
from sort import Sort
import cv2, numpy as np
model = YOLO('Potato_model_best_weights.pt')
tracker = Sort(max_age=10)
CountedIDs = set()

polygon_zone = np.array([[3, 183], [306, 45], [472, 209], [113, 373]], np.int32)
colors = [(255, 0, 0), (0, 255, 0), (0, 0, 255), (255, 255, 0), (255, 0, 255), (0, 255, 255), (255, 128, 0)]

cap = cv2.VideoCapture('ooo.mp4')
fps = int(cap.get(cv2.CAP_PROP_FPS))
width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
out = cv2.VideoWriter('output_khoai_tay.mp4', cv2.VideoWriter_fourcc(*'mp4v'), fps, (width, height))

while cap.isOpened():
    ret, frame = cap.read()
    if not ret: break
    result = next(model(frame, stream=True, device=0))
    
    dets = np.empty((0, 5))
    for box in result.boxes:
        x1, y1, x2, y2 = map(int, box.xyxy[0])
        dets = np.vstack((dets, np.array([x1, y1, x2, y2, box.conf[0].item()])))

    for det in tracker.update(dets):
        x1, y1, x2, y2, tid = map(int, det)
        color = colors[tid % 7]
        cv2.rectangle(frame, (x1, y1), (x2, y2), color, 2)
        cv2.putText(frame, str(tid), (x1, y1-5), cv2.FONT_HERSHEY_SIMPLEX, 0.5, color, 2)
        
        cx, cy = (x1 + x2) // 2, (y1 + y2) // 2
        if cv2.pointPolygonTest(polygon_zone, (cx, cy), False) >= 0:
            CountedIDs.add(tid)

    cv2.putText(frame, f'Counted IDs: {len(CountedIDs)}', (10, 40), cv2.FONT_HERSHEY_SIMPLEX, 1, (0, 255, 0), 2)
    cv2.polylines(frame, [polygon_zone], True, (0, 255, 0), 2)
    out.write(frame)

cap.release()
out.release()
print(f'Done! Total Counted: {len(CountedIDs)} Potatos.')
'''
    doc.add_paragraph(code_block)

    run = p.add_run('\nBước 4: Thành quả\n')
    run.bold = True
    p.add_run('Sau khi thấy dòng chữ [Done! Total Counted: XY Potatos.], bạn bấm Refresh tải lại mục Files bên trái. File "output_khoai_tay.mp4" sẽ xuất hiện. Tải về và mang chiếu lúc thuyết trình.')

    # PHẦN 4
    doc.add_heading('PHẦN 4: Hướng dẫn Train thêm vật thể mới (Fine-Tuning/Transfer Learning) - Dành cho Thành Viên 1 & 2', level=1)
    p = doc.add_paragraph('Đây là cách "chèn thêm" trí khôn nhận diện Quả Táo vào mạng Neural (file .pt) đang sẵn có mà không phải dạy lại từ đầu.\n')
    run = p.add_run('Bước 1: ')
    run.bold = True
    p.add_run('Lên công cụ Roboflow (Roboflow.com) xin 1 Dataset chứa Cả ảnh Khoai Tây và Ảnh Quả Táo.\n')
    run = p.add_run('Bước 2: ')
    run.bold = True
    p.add_run('Xuất (Export) Dataset đó dưới dạng định dạng YOLOv8 (Nó sẽ đẻ ra 1 file tên là data.yaml).\n')
    run = p.add_run('Bước 3: ')
    run.bold = True
    p.add_run('Lên Google Colab, đẩy cả bộ File Dataset nãy bạn kéo về kèm chung với con AI Cũ (Potato_model_best_weights.pt) lên.\n')
    run = p.add_run('Bước 4: ')
    run.bold = True
    p.add_run('Viết đoạn code sau ném vào Colab và chạy:\n')
    doc.add_paragraph('!yolo task=detect mode=train model=Potato_model_best_weights.pt data=data.yaml epochs=50 imgsz=640', style='Intense Quote')

    doc.save('HUONG_DAN_POTATO_COUNTER_TEAM.docx')
    print("File docx created successfully!")

if __name__ == "__main__":
    create_doc()
