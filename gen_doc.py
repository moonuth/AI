from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '4472C4')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def shade_cell(cell, color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    for side in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '6')
        b.set(qn('w:space'), '4')
        b.set(qn('w:color'), '4472C4')
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'EBF3FB')
    pPr.append(shd)
    return p

def add_note_box(doc, text, color='FFF2CC'):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    run = p.add_run('📌  ' + text)
    run.font.size = Pt(10)
    run.font.italic = True
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    pPr.append(shd)

def bold_run(para, text):
    r = para.add_run(text)
    r.bold = True
    return r

def create_full_doc():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════
    # TRANG BÌA
    # ══════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('🥔 HỆ THỐNG ĐẾM KHOAI TÂY TỰ ĐỘNG')
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run('POTATO COUNTER — TÀI LIỆU KỸ THUẬT TOÀN DIỆN')
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('Giải thích Kiến trúc · Thuật toán · Hướng dẫn Triển khai · Phân công Nhóm 5 Người\n')
    info.add_run('Dựa trên: YOLOv11 · SORT Tracker · Kalman Filter · Hungarian Algorithm\n')
    r3 = info.add_run('Tham chiếu lý thuyết: Russell & Norvig (AIMA 2021) · Pourret et al. (Bayesian Networks 2008)')
    r3.font.italic = True
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x59, 0x59, 0x59)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # MỤC LỤC (thủ công)
    # ══════════════════════════════════════════════════════════
    doc.add_heading('MỤC LỤC', level=1)
    toc_items = [
        ('PHẦN 1', 'Tổng quan dự án & Kiến trúc hệ thống'),
        ('PHẦN 2', 'Giải thích từng file trong dự án'),
        ('PHẦN 3', 'Thuật toán 1 — YOLOv11 (Phát hiện vật thể)'),
        ('PHẦN 4', 'Thuật toán 2 — Kalman Filter (Dự đoán vị trí)'),
        ('PHẦN 5', 'Thuật toán 3 — Hungarian Algorithm (Ghép detection & track)'),
        ('PHẦN 6', 'Thuật toán 4 — Polygon Zone (Đếm không trùng lặp)'),
        ('PHẦN 7', 'Liên hệ lý thuyết với 2 cuốn sách tham khảo'),
        ('PHẦN 8', 'Hướng dẫn cài đặt & chạy trên máy cá nhân (Local)'),
        ('PHẦN 9', 'Hướng dẫn deploy lên Google Colab (Cloud GPU)'),
        ('PHẦN 10', 'Hướng dẫn Train lại / Fine-Tuning mô hình'),
        ('PHẦN 11', 'Phân công nhóm 5 người & Git Workflow'),
        ('PHẦN 12', 'Bảng thông số quan trọng & Troubleshooting'),
    ]
    for num, title_text in toc_items:
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(2)
        bold_run(p, f'{num}: ')
        p.add_run(title_text)
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 1: TỔNG QUAN
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 1: Tổng quan dự án & Kiến trúc hệ thống', level=1)

    doc.add_heading('1.1 Mô tả bài toán', level=2)
    p = doc.add_paragraph(
        'Dự án xây dựng một hệ thống đếm khoai tây tự động hoạt động theo thời gian thực trên băng chuyền. '
        'Camera ghi lại video băng chuyền, phần mềm xử lý từng khung hình, phát hiện từng củ khoai, '
        'gán ID duy nhất cho mỗi củ, và đếm chúng khi đi qua vùng kiểm soát — đảm bảo không đếm trùng.'
    )

    doc.add_heading('1.2 Công nghệ sử dụng', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, txt in enumerate(['Thành phần', 'Công nghệ', 'Vai trò']):
        hdr[i].text = txt
        hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '4472C4')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    rows_data = [
        ('Phát hiện vật thể', 'YOLOv11m (custom)', 'Tìm bounding box khoai trong mỗi frame'),
        ('Theo dõi đối tượng', 'SORT + Kalman Filter', 'Gán ID nhất quán qua nhiều frame'),
        ('Tối ưu ghép cặp', 'Hungarian Algorithm', 'Ghép detection mới ↔ track đang có'),
        ('Xử lý video', 'OpenCV', 'Đọc video, vẽ kết quả, hiển thị'),
        ('Vùng đếm', 'Polygon Zone', 'Đếm khi tâm khoai vào vùng định sẵn'),
        ('Deep Learning backend', 'PyTorch + Ultralytics', 'Chạy mạng neural YOLO'),
    ]
    for rd in rows_data:
        row = tbl.add_row().cells
        for i, v in enumerate(rd):
            row[i].text = v

    doc.add_heading('1.3 Sơ đồ kiến trúc tổng thể', level=2)
    add_code_block(doc,
        "Video (potato.mp4)\n"
        "       │\n"
        "       ▼\n"
        " [OpenCV VideoCapture] ← đọc từng frame\n"
        "       │\n"
        "       ▼\n"
        " [Resize → 1280×720]   ← tối ưu tốc độ\n"
        "       │\n"
        "       ▼\n"
        " [YOLOv11 (imgsz=480)] ← phát hiện khoai → [x1,y1,x2,y2,conf] × N củ\n"
        "       │\n"
        "       ▼\n"
        " [SORT Tracker]\n"
        "    ├── Kalman Predict  ← dự đoán vị trí M track hiện có\n"
        "    ├── IOU Matrix      ← tính độ trùng lặp N×M\n"
        "    ├── Hungarian Algo  ← ghép tối ưu detection ↔ track\n"
        "    ├── Kalman Update   ← cập nhật track đã ghép\n"
        "    └── Quản lý track   ← tạo mới / xóa track cũ\n"
        "       │\n"
        "       ▼  [x1,y1,x2,y2,track_id] × M track\n"
        " [Polygon Zone Test]   ← cx,cy có trong vùng đa giác không?\n"
        "       │\n"
        "       ▼\n"
        " [CountedIDs (Set)]    ← thêm track_id → tự loại trùng\n"
        "       │\n"
        "       ▼\n"
        " [Hiển thị OpenCV]     ← vẽ box, ID, polygon, số đếm"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 2: FILE BREAKDOWN
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 2: Giải thích từng file trong dự án', level=1)

    files_info = [
        ('manual_counter.py', 'File chính — điểm khởi chạy chương trình',
         [
             ('Dòng 9–10', 'Xây đường dẫn tới file model (.pt) và video'),
             ('Dòng 13–14', 'Mở video bằng OpenCV, tạo cửa sổ hiển thị'),
             ('Dòng 16', 'Nạp model YOLO từ file weights vào bộ nhớ'),
             ('Dòng 19', 'Định nghĩa đa giác đếm (4 điểm tọa độ pixel tại độ phân giải 1280×720)'),
             ('Dòng 20', 'Tạo SORT tracker với max_age=10 (giữ track 10 frame khi mất)'),
             ('Dòng 21', 'CountedIDs là Set Python — tự loại bỏ trùng lặp'),
             ('Dòng 24–32', '7 màu BGR xoay vòng theo track_id để dễ phân biệt'),
             ('Dòng 35–86', 'Vòng lặp chính: đọc frame → YOLO → SORT → đếm → hiển thị'),
             ('Dòng 45', 'Resize mỗi frame về 1280×720 để tăng tốc độ xử lý'),
             ('Dòng 49–50', 'Chạy YOLO với stream=True và imgsz=480 (nhỏ hơn 640 → nhanh hơn ~30%)'),
             ('Dòng 53–57', 'Vòng lặp lấy tọa độ box + confidence, stack thành mảng numpy'),
             ('Dòng 60', 'Gửi detections vào SORT, nhận lại tracked objects có track_id'),
             ('Dòng 61–76', 'Vẽ box, số ID, kiểm tra điểm tâm có trong polygon không'),
             ('Dòng 79–82', 'Hiển thị số đếm tổng, vẽ polygon xanh lên frame'),
             ('Dòng 85', 'Nhấn Q để thoát'),
         ]),
        ('sort.py', 'Thuật toán SORT tracking (Alex Bewley, 2016)',
         [
             ('KalmanBoxTracker', 'Lớp đại diện 1 củ khoai đang được theo dõi — dùng Kalman Filter 7 chiều'),
             ('iou_batch()', 'Tính IOU giữa tất cả cặp detection/tracker cùng lúc (vectorized numpy)'),
             ('associate_detections_to_trackers()', 'Dùng Hungarian Algorithm ghép detection ↔ tracker tối ưu'),
             ('Sort.update()', 'Hàm gọi mỗi frame: predict → associate → update → quản lý vòng đời track'),
         ]),
        ('Potato_model_best_weights.pt', 'File trọng số mạng neural YOLOv11m đã huấn luyện',
         [
             ('Định dạng', 'PyTorch .pt — chứa toàn bộ weights của mạng CNN'),
             ('Kích thước', '~30–100 MB'),
             ('Lưu ý', 'KHÔNG commit file này lên Git (dùng .gitignore), chia sẻ qua Google Drive'),
         ]),
        ('requirements.txt', 'Danh sách thư viện Python cần cài',
         [
             ('ultralytics', 'Framework YOLO — load model, chạy inference'),
             ('opencv-python', 'Đọc video, xử lý ảnh, vẽ, pointPolygonTest'),
             ('filterpy', 'Thư viện Kalman Filter dùng trong sort.py'),
             ('lap', 'Hungarian Algorithm nhanh (Jonker-Volgenant)'),
             ('numpy', 'Tính toán ma trận bounding box'),
             ('torch / torchvision', 'Backend deep learning cho YOLO'),
             ('scikit-image', 'Dùng bởi sort.py cho image I/O'),
         ]),
        ('Inference_data/', 'Thư mục chứa video đầu vào',
         [
             ('potato.mp4', 'Video chính dùng bởi manual_counter.py'),
             ('ooo.mp4', 'Video dùng trong script Google Colab'),
             ('ppp.png', 'Ảnh frame mẫu để tham khảo'),
         ]),
        ('frame_1280x720.jpg', 'Ảnh frame mẫu tại độ phân giải 1280×720',
         [
             ('Mục đích', 'Upload lên Roboflow PolygonZone để vẽ lại vùng đếm'),
         ]),
    ]

    for fname, fdesc, fdetails in files_info:
        h = doc.add_heading(fname, level=2)
        doc.add_paragraph(fdesc)
        tbl = doc.add_table(rows=1, cols=2)
        set_table_border(tbl)
        hdr = tbl.rows[0].cells
        for i, txt in enumerate(['Vị trí / Thành phần', 'Chức năng']):
            hdr[i].text = txt
            hdr[i].paragraphs[0].runs[0].bold = True
            shade_cell(hdr[i], '2E75B6')
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        for loc, desc in fdetails:
            row = tbl.add_row().cells
            row[0].text = loc
            row[1].text = desc
        doc.add_paragraph()

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 3: YOLO
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 3: Thuật toán 1 — YOLOv11 (Phát hiện vật thể)', level=1)

    doc.add_heading('3.1 YOLO là gì?', level=2)
    doc.add_paragraph(
        'YOLO (You Only Look Once) là thuật toán phát hiện vật thể thời gian thực. Thay vì trượt cửa sổ '
        'qua ảnh nhiều lần như các phương pháp cũ (R-CNN, DPM), YOLO xử lý toàn bộ ảnh chỉ trong '
        'MỘT lần chạy qua mạng neural (1 forward pass), do đó cực kỳ nhanh.'
    )
    add_note_box(doc,
        'Liên hệ sách: Russell & Norvig (AIMA 2021) — Chương 21 "Deep Learning" mô tả '
        'CNN là nền tảng cho bài toán nhận diện ảnh. YOLO chính là ứng dụng CNN trong '
        'bài toán Object Detection.'
    )

    doc.add_heading('3.2 Cách YOLO xử lý 1 frame', level=2)
    add_code_block(doc,
        "Ảnh đầu vào (480×480 — sau khi resize)\n"
        "        │\n"
        "        ▼\n"
        "  [Backbone CNN]  ← Trích xuất đặc trưng (feature maps)\n"
        "        │\n"
        "        ▼\n"
        "  [Neck FPN]      ← Kết hợp đặc trưng nhiều tỉ lệ khác nhau\n"
        "        │\n"
        "        ▼\n"
        "  [Head]          ← Dự đoán cho từng ô lưới\n"
        "   Mỗi ô dự đoán:\n"
        "     - Bounding box: [x_center, y_center, width, height]\n"
        "     - Confidence score: P(có vật thể) × IoU(pred, truth)\n"
        "     - Class prob: P(potato | có vật thể)\n"
        "        │\n"
        "        ▼\n"
        "  [NMS — Non-Maximum Suppression]\n"
        "  Loại bỏ các box trùng, chỉ giữ box tốt nhất cho mỗi củ\n"
        "        │\n"
        "        ▼\n"
        "  Output: [[x1,y1,x2,y2, conf], ...]  ← mảng numpy gửi vào SORT"
    )

    doc.add_heading('3.3 Tham số tối ưu trong code', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Tham số', 'Giá trị dùng', 'Ý nghĩa']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('imgsz', '480 (mặc định 640)', 'Giảm độ phân giải YOLO xử lý → nhanh ~30% trên CPU/iGPU'),
        ('stream', 'True', 'Trả về generator thay vì list → tiết kiệm RAM'),
        ('frame resize', '1280×720', 'Thu nhỏ frame trước khi đưa vào → giảm tải xử lý'),
        ('device', '0 (GPU) hoặc cpu', 'Colab dùng device=0 để chạy trên T4 GPU'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 4: KALMAN FILTER
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 4: Thuật toán 2 — Kalman Filter (Dự đoán vị trí)', level=1)

    doc.add_heading('4.1 Tại sao cần Kalman Filter?', level=2)
    doc.add_paragraph(
        'Vấn đề thực tế: Khi 2 củ khoai đi gần nhau hoặc bị che khuất một vài frame, YOLO có thể '
        'không phát hiện được. Nếu không có cơ chế dự đoán, hệ thống sẽ mất track — củ khoai bị '
        'gán ID mới khi xuất hiện lại, dẫn đến ĐẾM TRÙNG. Kalman Filter giải quyết bằng cách '
        'DỰ ĐOÁN vị trí tiếp theo ngay cả khi không có quan sát mới.'
    )
    add_note_box(doc,
        'Liên hệ sách: Russell & Norvig (AIMA 2021) — Chương 4 "Search in Partially Observable '
        'Environments" và Chương 17 "Probabilistic Reasoning over Time". Kalman Filter là trường '
        'hợp đặc biệt của Bayes Filter khi mọi phân phối đều là Gaussian.\n\n'
        'Pourret et al. (Bayesian Networks 2008) — Chương về Dynamic Bayes Networks (DBN): '
        'Kalman Filter là một DBN đặc biệt với các biến ẩn liên tục và quan hệ tuyến tính.'
    )

    doc.add_heading('4.2 Vector trạng thái 7 chiều', level=2)
    add_code_block(doc,
        "x = [ cx,  cy,   s,    r,   vx,  vy,  vs ]\n"
        "       │    │    │     │     │    │    │\n"
        "       │    │    │     │     │    │    └── Tốc độ thay đổi diện tích (vận tốc s)\n"
        "       │    │    │     │     │    └─────── Tốc độ di chuyển theo trục Y\n"
        "       │    │    │     │     └──────────── Tốc độ di chuyển theo trục X\n"
        "       │    │    │     └────────────────── Tỉ lệ khung hình (width / height)\n"
        "       │    │    └──────────────────────── Diện tích bounding box (scale = w×h)\n"
        "       │    └───────────────────────────── Tọa độ Y tâm bounding box\n"
        "       └────────────────────────────────── Tọa độ X tâm bounding box\n\n"
        "Quan sát được (dim_z=4): cx, cy, s, r\n"
        "KHÔNG quan sát được:     vx, vy, vs  (Kalman tự ước lượng từ chuyển động)"
    )

    doc.add_heading('4.3 Hai bước chính của Kalman Filter', level=2)
    doc.add_paragraph('Mỗi frame, Kalman thực hiện 2 bước:')

    doc.add_heading('Bước 1 — PREDICT (Dự đoán)', level=3)
    add_code_block(doc,
        "# 'Củ khoai sẽ ở đâu frame tiếp theo, dựa trên vận tốc hiện tại?'\n\n"
        "x̂ = F · x          # Nhân ma trận chuyển trạng thái (mô hình vận tốc không đổi)\n"
        "P̂ = F · P · Fᵀ + Q  # Độ không chắc chắn tăng lên (Q = nhiễu quá trình)\n\n"
        "# Ma trận F (7×7) — mô hình chuyển động thẳng đều:\n"
        "# [1 0 0 0 1 0 0]   cx_mới  = cx_cũ  + vx\n"
        "# [0 1 0 0 0 1 0]   cy_mới  = cy_cũ  + vy\n"
        "# [0 0 1 0 0 0 1]   s_mới   = s_cũ   + vs\n"
        "# [0 0 0 1 0 0 0]   r_mới   = r_cũ   (không đổi)\n"
        "# [0 0 0 0 1 0 0]   vx_mới  = vx_cũ\n"
        "# [0 0 0 0 0 1 0]   vy_mới  = vy_cũ\n"
        "# [0 0 0 0 0 0 1]   vs_mới  = vs_cũ"
    )

    doc.add_heading('Bước 2 — UPDATE (Cập nhật khi có quan sát từ YOLO)', level=3)
    add_code_block(doc,
        "# 'Kết hợp dự đoán + bounding box thực tế YOLO phát hiện'\n\n"
        "K = P̂ · Hᵀ · (H · P̂ · Hᵀ + R)⁻¹   # Kalman Gain\n"
        "x = x̂ + K · (z - H · x̂)             # Cập nhật trạng thái\n"
        "P = (I - K · H) · P̂                  # Giảm độ không chắc chắn\n\n"
        "# z = quan sát thực từ YOLO: [cx_yolo, cy_yolo, s_yolo, r_yolo]\n"
        "# H = ma trận quan sát (4×7): chiếu từ không gian 7D xuống 4D\n"
        "# R = ma trận nhiễu đo lường (YOLO không hoàn toàn chính xác)\n\n"
        "# Kalman Gain K điều chỉnh:\n"
        "#   K gần 1 → tin quan sát hơn dự đoán\n"
        "#   K gần 0 → tin dự đoán hơn quan sát (YOLO đang không chính xác)"
    )

    doc.add_heading('4.4 Tham số Kalman trong code', level=2)
    add_code_block(doc,
        "self.kf.R[2:,2:]   *= 10.     # Tăng nhiễu đo lường cho s, r → ít tin YOLO về size\n"
        "self.kf.P[4:,4:]   *= 1000.   # Uncertainty cao cho vận tốc ban đầu (chưa biết)\n"
        "self.kf.P          *= 10.\n"
        "self.kf.Q[-1,-1]   *= 0.01    # Giảm nhiễu quá trình cho vs\n"
        "self.kf.Q[4:,4:]   *= 0.01    # Giảm nhiễu quá trình cho vx, vy, vs\n\n"
        "Sort(max_age=10)              # Giữ track 10 frame dù không thấy khoai"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 5: HUNGARIAN
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 5: Thuật toán 3 — Hungarian Algorithm (Ghép detection & track)', level=1)

    doc.add_heading('5.1 Bài toán đặt ra', level=2)
    doc.add_paragraph(
        'Mỗi frame, YOLO cho ra N bounding box mới. SORT đang theo dõi M track cũ. '
        'Câu hỏi: box nào tương ứng với track nào? Đây là bài toán gán nhãn tối ưu '
        '(Assignment Problem) — tìm cách ghép 1-1 sao cho tổng "chi phí" (1 - IOU) nhỏ nhất.'
    )
    add_note_box(doc,
        'Liên hệ sách: Russell & Norvig (AIMA 2021) — Chương 17 "Making Complex Decisions" '
        'và Optimization. Bài toán Hungarian là bài toán tối ưu hóa tổ hợp điển hình trong AI, '
        'được giải bằng thuật toán thời gian đa thức O(n³).'
    )

    doc.add_heading('5.2 Bước 1 — Tính ma trận IOU', level=2)
    doc.add_paragraph(
        'IOU (Intersection over Union) đo mức độ trùng khớp giữa 2 bounding box:'
    )
    add_code_block(doc,
        "IOU = Diện_tích_giao / Diện_tích_hợp\n\n"
        "      ┌──────┐\n"
        "      │  A   │\n"
        "      │  ┌───┼──┐\n"
        "      │  │///│  │   /// = phần giao nhau\n"
        "      └──┼───┘  │\n"
        "         │  B   │\n"
        "         └──────┘\n\n"
        "IOU = 0.0  → 2 box không liên quan\n"
        "IOU = 1.0  → 2 box hoàn toàn trùng nhau\n\n"
        "Ví dụ ma trận IOU (3 track × 4 detections):\n"
        "             Det_A  Det_B  Det_C  Det_D\n"
        "  Track_1  [  0.8    0.1    0.0    0.0  ]\n"
        "  Track_2  [  0.1    0.9    0.0    0.0  ]\n"
        "  Track_3  [  0.0    0.1    0.7    0.1  ]\n"
        "  → Det_D không khớp track nào → tạo Track_4 mới!"
    )

    doc.add_heading('5.3 Bước 2 — Hungarian Algorithm tìm ghép tối ưu', level=2)
    add_code_block(doc,
        "# Mục tiêu: Maximize tổng IOU của tất cả cặp được ghép\n"
        "# (Tương đương: Minimize tổng chi phí = 1 - IOU)\n\n"
        "Kết quả tối ưu:\n"
        "  Det_A → Track_1  (IOU = 0.8) ✓\n"
        "  Det_B → Track_2  (IOU = 0.9) ✓\n"
        "  Det_C → Track_3  (IOU = 0.7) ✓\n"
        "  Det_D → Track_4  (MỚI, không khớp track cũ nào)\n\n"
        "Trong code dùng thư viện LAP (Jonker-Volgenant — biến thể nhanh hơn Hungarian gốc):\n\n"
        "def linear_assignment(cost_matrix):\n"
        "    import lap\n"
        "    _, x, y = lap.lapjv(cost_matrix, extend_cost=True)\n"
        "    return np.array([[y[i], i] for i in x if i >= 0])"
    )

    doc.add_heading('5.4 Xử lý sau khi ghép', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Trường hợp', 'Hành động']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Detection ghép được track (IOU ≥ 0.3)', 'Cập nhật Kalman Filter của track đó bằng vị trí mới'),
        ('Detection KHÔNG khớp track nào (IOU < 0.3)', 'Tạo KalmanBoxTracker mới → track mới với ID mới'),
        ('Track KHÔNG khớp detection nào', 'Tăng time_since_update; nếu > max_age → xóa track'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_paragraph()
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 6: POLYGON ZONE
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 6: Thuật toán 4 — Polygon Zone (Đếm không trùng lặp)', level=1)

    doc.add_heading('6.1 Logic đếm', level=2)
    add_code_block(doc,
        "# Tính tâm bounding box:\n"
        "cx = (x1 + x2) // 2\n"
        "cy = (y1 + y2) // 2\n\n"
        "# Kiểm tra tâm có nằm trong đa giác đếm không:\n"
        "if cv2.pointPolygonTest(polygon_zone, (cx, cy), False) >= 0:\n"
        "    CountedIDs.add(track_id)  # Set tự loại bỏ trùng lặp!\n\n"
        "# pointPolygonTest trả về:\n"
        "#  +1.0 → điểm nằm TRONG đa giác\n"
        "#   0.0 → điểm nằm TRÊN cạnh đa giác\n"
        "#  -1.0 → điểm nằm NGOÀI đa giác\n"
        "# Điều kiện >= 0 bắt cả trong và trên cạnh"
    )

    doc.add_heading('6.2 Thuật toán Ray Casting (bên trong pointPolygonTest)', level=2)
    doc.add_paragraph(
        'OpenCV sử dụng thuật toán Ray Casting để kiểm tra điểm trong đa giác:'
    )
    add_code_block(doc,
        "Từ điểm (cx, cy), chiếu 1 tia ra vô cực theo chiều ngang.\n"
        "Đếm số lần tia đó cắt các cạnh của đa giác.\n\n"
        "  Số lần cắt LẺ  → điểm NẰM TRONG đa giác ✓\n"
        "  Số lần cắt CHẴN → điểm NẰM NGOÀI đa giác ✗\n\n"
        "Ví dụ với polygon 4 điểm: [[3,183],[306,45],[472,209],[113,373]]\n"
        "  → Hình thang xiên bao phủ băng chuyền bên TRÁI\n"
        "  → Tọa độ tính tại độ phân giải 1280×720"
    )

    doc.add_heading('6.3 Tại sao dùng Set thay vì biến đếm thông thường?', level=2)
    add_code_block(doc,
        "# ❌ Sai — đếm biến thông thường:\n"
        "count += 1  # Khoai đi qua vùng 10 frame → đếm 10 lần!\n\n"
        "# ✅ Đúng — dùng Set:\n"
        "CountedIDs = set()\n"
        "CountedIDs.add(5)   # {5}\n"
        "CountedIDs.add(5)   # {5}  ← không thay đổi, ID=5 đã có\n"
        "CountedIDs.add(7)   # {5, 7}\n"
        "len(CountedIDs)     # 2  ← đúng, chỉ 2 củ khoai khác nhau\n\n"
        "# Nhờ Kalman Filter + Hungarian: cùng 1 củ khoai luôn có cùng track_id\n"
        "# Nhờ Set: track_id đó chỉ được đếm 1 lần dù vào vùng 100 frame"
    )

    doc.add_heading('6.4 Cách thay đổi vùng đếm', level=2)
    p = doc.add_paragraph()
    bold_run(p, 'Để vẽ lại polygon cho băng chuyền khác:')
    steps = [
        'Lấy file frame_1280x720.jpg (đã có sẵn trong thư mục dự án)',
        'Truy cập https://polygonzone.roboflow.com/',
        'Upload ảnh lên → nhấn phím P → click các góc của vùng cần đếm',
        'Roboflow xuất ra mảng numpy tọa độ',
        'Dán vào dòng 19 của manual_counter.py:',
    ]
    for i, s in enumerate(steps, 1):
        doc.add_paragraph(f'  {i}. {s}')
    add_code_block(doc, "polygon_zone = np.array([[x1,y1],[x2,y2],[x3,y3],[x4,y4]], np.int32)")
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 7: LÝ THUYẾT TỪ SÁCH
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 7: Liên hệ lý thuyết với 2 cuốn sách tham khảo', level=1)

    doc.add_heading('7.1 Russell & Norvig — Artificial Intelligence: A Modern Approach (2021)', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Chương trong AIMA', 'Khái niệm', 'Ứng dụng trong Potato Counter']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '1F497D')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Chương 4 — Uncertain Knowledge', 'Bayesian Filtering, Hidden Markov Model', 'Nền tảng lý thuyết của Kalman Filter trong SORT'),
        ('Chương 17 — Probabilistic Reasoning over Time', 'Kalman Filter, State Estimation', 'KalmanBoxTracker trong sort.py: predict & update 2 bước'),
        ('Chương 21 — Deep Learning', 'CNN, Feature Maps, Backpropagation', 'Backbone của YOLOv11 dùng CNN để trích xuất đặc trưng'),
        ('Chương 21 — Object Detection', 'Bounding Box Regression, NMS, Anchor Boxes', 'Cách YOLOv11 output [x1,y1,x2,y2,conf]'),
        ('Chương 17 — Optimization', 'Assignment Problem, Linear Programming', 'Hungarian Algorithm trong associate_detections_to_trackers()'),
        ('Chương 25 — Computer Vision', 'Image features, Tracking', 'Toàn bộ pipeline detection + tracking'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_paragraph()

    doc.add_heading('7.2 Pourret, Naïm, Marcot — Bayesian Networks: A Practical Guide (2008)', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Chương trong sách', 'Khái niệm', 'Ứng dụng trong Potato Counter']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '1F497D')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Chương về Inference', 'Belief Propagation, P(x|evidence)', 'Kalman Update step: cập nhật P(vị_trí | quan_sát_YOLO)'),
        ('Chương về Dynamic BN', 'Temporal Models, DBN', 'SORT là một Dynamic Bayes Network: trạng thái t phụ thuộc t-1'),
        ('Chương về Uncertainty', 'Prior, Likelihood, Posterior (Bayes Theorem)', 'Kalman Gain K cân bằng giữa prior (dự đoán) và likelihood (YOLO)'),
        ('Chương về Learning', 'Parameter Learning từ dữ liệu', 'Huấn luyện YOLOv11 = học tham số mạng từ ảnh có nhãn'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_paragraph()

    doc.add_heading('7.3 Tóm tắt liên hệ toàn bộ', level=2)
    add_code_block(doc,
        "Bayesian Framework (Pourret et al.)\n"
        "  └── P(trạng_thái | quan_sát) = P(quan_sát | trạng_thái) × P(trạng_thái)\n"
        "                                  ───────────────────────────────────────────\n"
        "                                          P(quan_sát)\n\n"
        "Trong Kalman Filter (AIMA Ch.17):\n"
        "  Prior    = Kalman Predict   → x̂ = F·x  (dự đoán từ frame trước)\n"
        "  Likelihood = YOLO detection  → z = [cx, cy, s, r]  (quan sát mới)\n"
        "  Posterior = Kalman Update    → x = x̂ + K(z - H·x̂)  (kết hợp tốt nhất)\n\n"
        "Kalman Gain K = 'bao nhiêu % tin YOLO vs tin dự đoán'\n"
        "  K → 1: tin YOLO hoàn toàn (quan sát rõ ràng)\n"
        "  K → 0: tin dự đoán hoàn toàn (YOLO không chính xác)"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 8: CÀI ĐẶT LOCAL
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 8: Hướng dẫn cài đặt & chạy trên máy cá nhân (Local)', level=1)

    doc.add_heading('8.1 Yêu cầu hệ thống', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Thành phần', 'Yêu cầu']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Python', '3.9 – 3.12 (KHÔNG dùng 3.13 — PyTorch chưa hỗ trợ)'),
        ('RAM', 'Tối thiểu 4 GB (khuyến nghị 8 GB)'),
        ('Ổ cứng', '~3 GB trống cho PyTorch + model'),
        ('Hệ điều hành', 'Windows 10/11, macOS, Ubuntu'),
        ('GPU (không bắt buộc)', 'NVIDIA CUDA để chạy nhanh hơn'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v

    doc.add_heading('8.2 Các bước cài đặt', level=2)
    steps_install = [
        ('Bước 1 — Cài Python', 'Tải từ python.org, chọn bản 3.9–3.12.\nWindows: ✅ TICK vào "Add Python to PATH" trước khi nhấn Install.'),
        ('Bước 2 — Lấy code dự án',
         'git clone https://github.com/NHOM_BAN/Potato-Counter.git\ncd Potato-Counter'),
        ('Bước 3 — Tạo môi trường ảo (khuyến nghị)',
         'python -m venv .venv\n# Kích hoạt trên Windows:\n.venv\\Scripts\\activate\n# Kích hoạt trên Mac/Linux:\nsource .venv/bin/activate'),
        ('Bước 4 — Cài thư viện',
         'pip install -r requirements.txt\n# Chờ 5–15 phút, PyTorch ~700MB'),
        ('Bước 5 — Đặt file model',
         'Tải Potato_model_best_weights.pt từ Google Drive nhóm\nĐặt vào thư mục gốc Potato-Counter/'),
        ('Bước 6 — Chạy chương trình',
         'python manual_counter.py'),
    ]
    for title, code in steps_install:
        p = doc.add_paragraph()
        bold_run(p, title)
        add_code_block(doc, code)

    doc.add_heading('8.3 Điều khiển khi chạy', level=2)
    tbl = doc.add_table(rows=1, cols=2)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Phím / Hành động', 'Kết quả']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Q', 'Thoát chương trình'),
        ('Kéo góc cửa sổ', 'Thay đổi kích thước cửa sổ (WINDOW_NORMAL)'),
        ('Video kết thúc', 'Chương trình tự thoát, in kết quả ra terminal'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v

    doc.add_heading('8.4 Hiệu suất thực tế', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Phần cứng', 'FPS ước tính', 'Ghi chú']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('Intel Core i5 (không GPU)', '3–6 FPS', 'Chậm nhưng vẫn chạy được, dùng cho test'),
        ('Intel Core i7/i9 + iGPU', '6–12 FPS', 'Dùng được cho demo nhỏ'),
        ('NVIDIA GTX 1060+', '20–35 FPS', 'Cần cài CUDA + PyTorch GPU'),
        ('Google Colab T4 GPU', '30–50 FPS', 'Lý tưởng cho demo và thuyết trình'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 9: COLAB
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 9: Hướng dẫn deploy lên Google Colab (Cloud GPU)', level=1)

    doc.add_heading('9.1 Các bước thực hiện', level=2)
    colab_steps = [
        ('Bước 1 — Mở Colab', 'Vào colab.research.google.com → Tạo Notebook mới'),
        ('Bước 2 — Bật GPU',
         'Runtime → Change runtime type → T4 GPU → Save\n(Miễn phí, giới hạn ~12h/ngày)'),
        ('Bước 3 — Upload 3 file cốt lõi',
         'Click biểu tượng 📁 bên trái → Upload:\n  - Potato_model_best_weights.pt\n  - sort.py\n  - ooo.mp4  (hoặc potato.mp4)'),
        ('Bước 4 — Dán code vào cell', 'Xem script đầy đủ bên dưới'),
        ('Bước 5 — Tải kết quả về', 'Sau khi xong, refresh Files panel → download output_khoai_tay.mp4'),
    ]
    for title, detail in colab_steps:
        p = doc.add_paragraph()
        bold_run(p, title + ': ')
        p.add_run(detail)

    doc.add_heading('9.2 Script Colab đầy đủ', level=2)
    add_code_block(doc,
        "# Cell 1: Cài thư viện\n"
        "!pip install ultralytics scikit-image filterpy lap\n\n"
        "# Cell 2: Chạy hệ thống đếm (xuất ra video)\n"
        "from ultralytics import YOLO\n"
        "from sort import Sort\n"
        "import cv2, numpy as np\n\n"
        "model = YOLO('Potato_model_best_weights.pt')\n"
        "tracker = Sort(max_age=10)\n"
        "CountedIDs = set()\n"
        "polygon_zone = np.array([[3,183],[306,45],[472,209],[113,373]], np.int32)\n"
        "colors = [(255,0,0),(0,255,0),(0,0,255),(255,255,0),(255,0,255),(0,255,255),(255,128,0)]\n\n"
        "cap = cv2.VideoCapture('ooo.mp4')\n"
        "fps    = int(cap.get(cv2.CAP_PROP_FPS))\n"
        "width  = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))\n"
        "height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))\n"
        "out = cv2.VideoWriter('output_khoai_tay.mp4',\n"
        "                      cv2.VideoWriter_fourcc(*'mp4v'), fps, (width, height))\n\n"
        "while cap.isOpened():\n"
        "    ret, frame = cap.read()\n"
        "    if not ret: break\n"
        "    result = next(model(frame, stream=True, device=0))  # device=0 = GPU\n"
        "    dets = np.empty((0, 5))\n"
        "    for box in result.boxes:\n"
        "        x1,y1,x2,y2 = map(int, box.xyxy[0])\n"
        "        dets = np.vstack((dets, [x1,y1,x2,y2,box.conf[0].item()]))\n"
        "    for det in tracker.update(dets):\n"
        "        x1,y1,x2,y2,tid = map(int, det)\n"
        "        color = colors[tid % 7]\n"
        "        cv2.rectangle(frame,(x1,y1),(x2,y2),color,2)\n"
        "        cv2.putText(frame,str(tid),(x1,y1-5),cv2.FONT_HERSHEY_SIMPLEX,0.5,color,2)\n"
        "        cx,cy = (x1+x2)//2,(y1+y2)//2\n"
        "        if cv2.pointPolygonTest(polygon_zone,(cx,cy),False) >= 0:\n"
        "            CountedIDs.add(tid)\n"
        "    cv2.putText(frame,f'Counted: {len(CountedIDs)}',(10,40),\n"
        "                cv2.FONT_HERSHEY_SIMPLEX,1,(0,255,0),2)\n"
        "    cv2.polylines(frame,[polygon_zone],True,(0,255,0),2)\n"
        "    out.write(frame)\n\n"
        "cap.release(); out.release()\n"
        "print(f'✅ Xong! Tổng số khoai đếm được: {len(CountedIDs)} củ')"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 10: TRAIN LẠI
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 10: Hướng dẫn Train lại / Fine-Tuning mô hình', level=1)

    doc.add_heading('10.1 Khi nào cần train lại?', level=2)
    situations = [
        'Model nhận diện sai nhiều (độ chính xác thấp)',
        'Muốn thêm loại vật thể mới (vd: quả táo, cà rốt)',
        'Video từ camera góc khác, ánh sáng khác',
        'Băng chuyền khác với dataset gốc',
    ]
    for s in situations:
        doc.add_paragraph(f'• {s}')

    doc.add_heading('10.2 Các bước Fine-Tuning trên Google Colab', level=2)
    add_code_block(doc,
        "# Bước 1: Lấy dataset trên Roboflow\n"
        "# Vào roboflow.com → tìm dataset khoai tây (hoặc tự annotate)\n"
        "# Export → Format: YOLOv8 → Download ZIP\n"
        "# ZIP chứa: images/train/, images/val/, labels/..., data.yaml\n\n"
        "# Bước 2: Upload lên Colab\n"
        "# Upload: Potato_model_best_weights.pt + toàn bộ thư mục dataset\n\n"
        "# Bước 3: Chạy lệnh train\n"
        "!yolo task=detect mode=train \\\n"
        "      model=Potato_model_best_weights.pt \\\n"
        "      data=data.yaml \\\n"
        "      epochs=50 \\\n"
        "      imgsz=640 \\\n"
        "      batch=16\n\n"
        "# Bước 4: Lấy model mới\n"
        "# Model được lưu tại: runs/detect/train/weights/best.pt\n"
        "# Download về, đổi tên, chia sẻ với nhóm qua Google Drive"
    )

    doc.add_heading('10.3 Giải thích tham số train', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Tham số', 'Giá trị gợi ý', 'Ý nghĩa']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    for r in [
        ('model', 'Potato_model_best_weights.pt', 'Bắt đầu từ model đã train (Transfer Learning)'),
        ('data', 'data.yaml', 'File config dataset: đường dẫn ảnh + tên lớp'),
        ('epochs', '50', 'Số lần duyệt qua toàn bộ dataset'),
        ('imgsz', '640', 'Độ phân giải ảnh khi train (lớn hơn → chính xác hơn nhưng chậm)'),
        ('batch', '16', 'Số ảnh xử lý cùng lúc (giảm nếu hết RAM GPU)'),
    ]:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 11: NHÓM & GIT
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 11: Phân công nhóm 5 người & Git Workflow', level=1)

    doc.add_heading('11.1 Phân công vai trò', level=2)
    tbl = doc.add_table(rows=1, cols=4)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Thành viên', 'Vai trò', 'Nhiệm vụ chính', 'File liên quan']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '1F497D')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    members = [
        ('Thành viên 1', 'AI Trainer', 'Thu thập dữ liệu, huấn luyện model YOLO mới, export .pt', 'Potato_model_best_weights.pt\ndata.yaml'),
        ('Thành viên 2', 'Data Annotator', 'Annotate ảnh trên Roboflow, kiểm tra chất lượng nhãn, validate độ chính xác', 'Dataset trên Roboflow'),
        ('Thành viên 3', 'Local Developer', 'Chạy manual_counter.py, điều chỉnh polygon, tối ưu tham số, sửa bugs giao diện', 'manual_counter.py\nframe_1280x720.jpg'),
        ('Thành viên 4', 'QA Tester', 'Test với nhiều video khác nhau, đo độ chính xác đếm, viết báo cáo lỗi', 'Inference_data/\nmanual_counter.py'),
        ('Thành viên 5', 'Cloud Deployment', 'Deploy lên Colab GPU, tạo video demo xuất ra, đo FPS, chuẩn bị thuyết trình', 'Colab notebook\noutput_khoai_tay.mp4'),
    ]
    for m in members:
        row = tbl.add_row().cells
        for i,v in enumerate(m): row[i].text = v

    doc.add_heading('11.2 Cấu trúc nhánh Git', level=2)
    add_code_block(doc,
        "main                         ← Nhánh chính, chỉ merge code ổn định đã test\n"
        "  └── dev                    ← Nhánh tích hợp, merge vào đây trước\n"
        "        ├── feat/tv1-retrain  ← Thành viên 1: thực nghiệm train model\n"
        "        ├── feat/tv2-dataset  ← Thành viên 2: chuẩn bị dataset\n"
        "        ├── feat/tv3-polygon  ← Thành viên 3: điều chỉnh polygon/UI\n"
        "        ├── feat/tv4-testing  ← Thành viên 4: script test, báo cáo\n"
        "        └── feat/tv5-colab    ← Thành viên 5: notebook Colab"
    )

    doc.add_heading('11.3 Quy trình làm việc hàng ngày', level=2)
    add_code_block(doc,
        "# 1. Luôn bắt đầu bằng pull code mới nhất từ dev\n"
        "git checkout dev\n"
        "git pull origin dev\n\n"
        "# 2. Chuyển sang nhánh của mình\n"
        "git checkout feat/tv3-polygon\n\n"
        "# 3. Làm việc, sửa code...\n\n"
        "# 4. Commit với message rõ ràng\n"
        "git add manual_counter.py\n"
        'git commit -m "Điều chỉnh polygon sang băng chuyền bên phải"\n\n'
        "# 5. Push nhánh của mình\n"
        "git push origin feat/tv3-polygon\n\n"
        "# 6. Tạo Pull Request trên GitHub → merge vào dev\n"
        "# 7. Sau khi cả nhóm review OK → merge dev vào main"
    )

    doc.add_heading('11.4 File .gitignore — QUAN TRỌNG', level=2)
    add_code_block(doc,
        "# Thêm vào file .gitignore:\n\n"
        "*.pt          # Model weights quá nặng (>30MB)\n"
        "*.mp4         # Video đầu vào/ra\n"
        "*.avi\n"
        "__pycache__/\n"
        ".venv/\n"
        "Inference_data/\n"
        "runs/          # Output của quá trình train YOLO\n\n"
        "# Chia sẻ file .pt và video qua Google Drive:\n"
        "# https://drive.google.com/drive/folders/LINK_NHOM"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════════════════════
    # PHẦN 12: THÔNG SỐ & TROUBLESHOOTING
    # ══════════════════════════════════════════════════════════
    doc.add_heading('PHẦN 12: Bảng thông số quan trọng & Troubleshooting', level=1)

    doc.add_heading('12.1 Bảng thông số cần biết', level=2)
    tbl = doc.add_table(rows=1, cols=4)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Thông số', 'Vị trí trong code', 'Giá trị hiện tại', 'Khi nào cần đổi']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], '2E75B6')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    params = [
        ('polygon_zone', 'manual_counter.py dòng 19', '4 điểm tại 1280×720', 'Đổi góc camera / băng chuyền khác'),
        ('max_age', 'Sort(max_age=10)', '10 frames', 'Tăng nếu khoai hay bị mất track; giảm nếu RAM ít'),
        ('imgsz', 'model(..., imgsz=480)', '480px', 'Giảm xuống 320 nếu quá chậm; tăng 640 nếu cần chính xác'),
        ('min_hits', 'Sort(min_hits=3)', '3 frames', 'Giảm xuống 1 nếu khoai bị mất quá sớm'),
        ('iou_threshold', 'Sort(iou_threshold=0.3)', '0.3', 'Tăng nếu đếm nhầm; giảm nếu bị mất track'),
        ('frame resize', 'cv2.resize(..., (1280,720))', '1280×720', 'Giảm xuống 640×360 nếu CPU quá chậm'),
    ]
    for r in params:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v

    doc.add_heading('12.2 Bảng xử lý lỗi thường gặp', level=2)
    tbl = doc.add_table(rows=1, cols=3)
    set_table_border(tbl)
    hdr = tbl.rows[0].cells
    for i, t in enumerate(['Lỗi / Vấn đề', 'Nguyên nhân', 'Cách sửa']):
        hdr[i].text = t; hdr[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr[i], 'C00000')
        hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
    errors = [
        ("No module named 'ultralytics'", 'Chưa cài thư viện', 'pip install -r requirements.txt'),
        ('Video frame is empty ngay lập tức', 'Sai đường dẫn video', 'Kiểm tra file potato.mp4 có trong Inference_data/'),
        ('No module named lap', 'lap chưa cài', 'pip install lap'),
        ('FPS 1–2, rất chậm', 'CPU yếu, không có GPU', 'Giảm imgsz=320, hoặc dùng Google Colab'),
        ('Khoai đếm trùng nhiều', 'max_age quá nhỏ, track bị mất', 'Tăng Sort(max_age=20) hoặc max_age=30'),
        ('Khoai đi qua không đếm', 'Polygon sai vị trí', 'Vẽ lại polygon tại polygonzone.roboflow.com'),
        ('Cửa sổ không hiển thị (Linux)', 'OpenCV không có DISPLAY', 'Dùng Colab (xuất video) thay vì hiển thị live'),
        ('CUDA out of memory (Colab)', 'GPU RAM đầy', 'Giảm imgsz=320 hoặc batch nhỏ hơn'),
        ('File .pt không tìm thấy', 'Không commit lên Git (đúng!)', 'Tải từ Google Drive nhóm, đặt vào thư mục gốc'),
        ('Màu box không nhất quán', 'track_id thay đổi mỗi lần chạy', 'Bình thường — ID gán theo thứ tự phát hiện'),
    ]
    for r in errors:
        row = tbl.add_row().cells
        for i,v in enumerate(r): row[i].text = v

    doc.add_heading('12.3 Quick Reference — Lệnh chạy nhanh', level=2)
    add_code_block(doc,
        "# === CÀI ĐẶT ===\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate          # Windows\n"
        "pip install -r requirements.txt\n\n"
        "# === CHẠY CHÍNH ===\n"
        "python manual_counter.py\n\n"
        "# === THAY ĐỔI THAM SỐ NHANH ===\n"
        "# Dòng 19: polygon_zone  → vùng đếm\n"
        "# Dòng 20: Sort(max_age=10) → thời gian giữ track\n"
        "# Dòng 45: cv2.resize(frame, (1280, 720)) → tốc độ\n"
        "# Dòng 49: imgsz=480 → độ phân giải YOLO\n\n"
        "# === PHÍM TẮT KHI CHẠY ===\n"
        "Q → Thoát chương trình"
    )

    # TRANG KẾT
    doc.add_page_break()
    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = end.add_run('― Hết tài liệu ―')
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        'Tài liệu này được tổng hợp dựa trên:\n'
        'Source code: manual_counter.py · sort.py\n'
        'Tham chiếu: Russell & Norvig — AI: A Modern Approach (2021) · '
        'Pourret et al. — Bayesian Networks (2008)'
    ).font.italic = True

    out_path = os.path.join(os.getcwd(), 'POTATO_COUNTER_TAI_LIEU_KY_THUAT.docx')
    doc.save(out_path)
    print(f'✅ Đã tạo file: {out_path}')

if __name__ == '__main__':
    create_full_doc()
